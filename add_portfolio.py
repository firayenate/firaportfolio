from docx import Document

def update_portfolio_link(file_path):
    doc = Document(file_path)

    # Check existing paragraphs for portfolio line and update
    found = False
    for para in doc.paragraphs:
        if "Portfolio:" in para.text:
            para.text = "Portfolio: https://firayenate.github.io/firaportfolio"
            found = True
            break

    if not found:
        doc.add_paragraph("Portfolio: https://firayenate.github.io/firaportfolio")

    doc.save(file_path)
    print("Done! Portfolio link updated in CV.")

if __name__ == "__main__":
    update_portfolio_link("Firaol_Desalegn_CV.docx")
