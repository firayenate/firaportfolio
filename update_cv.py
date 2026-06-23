from docx import Document

def update_cv(file_path):
    doc = Document(file_path)
    
    # [6] Programming Languages: Python (actively advancing daily), Java, C++
    if len(doc.paragraphs) > 6:
        doc.paragraphs[6].text = "Programming Languages: Python (Intermediate - 50%), Java, C++"
    
    # [7] Web Development: HTML, CSS, JavaScript (actively improving)
    if len(doc.paragraphs) > 7:
        doc.paragraphs[7].text = "Web Development: HTML, CSS, JavaScript (Intermediate - 50%), Vue.js (Beginner)"
        
    # [18] Personal Portfolio Website -> Rock Paper Scissors
    if len(doc.paragraphs) > 18:
        doc.paragraphs[18].text = "Rock Paper Scissors Game"
    if len(doc.paragraphs) > 19:
        doc.paragraphs[19].text = "A JavaScript game where you can play Rock, Paper, Scissors against the computer. Includes score tracking and interactive web elements."

    # [20] Shoe Shop -> Hotel Management System
    if len(doc.paragraphs) > 20:
        doc.paragraphs[20].text = "Hotel Management System"
    if len(doc.paragraphs) > 21:
        doc.paragraphs[21].text = "Developed a C++ console application to manage hotel room bookings, process check-outs, and track real-time room availability."

    # [22] Hotel Website (In Progress) -> Remove
    if len(doc.paragraphs) > 22:
        doc.paragraphs[22].clear()
    if len(doc.paragraphs) > 23:
        doc.paragraphs[23].clear()

    # Add GitHub to the end
    doc.add_paragraph("GitHub: https://github.com/firayenate")

    doc.save("Firaol_Desalegn_CV_Updated.docx")
    print("Updated CV saved to Firaol_Desalegn_CV_Updated.docx")

if __name__ == "__main__":
    update_cv("Firaol_Desalegn_CV.docx")
